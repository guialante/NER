"""
LLM integration for enhanced entity extraction and document generation from legal documents.
"""
import os
import json
import docx
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from typing import List, Dict, Optional
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

# Load environment variables from .env file
load_dotenv()

# Ensure API key is available
if not os.getenv("OPENAI_API_KEY"):
    raise ValueError("OPENAI_API_KEY not found in environment variables. Please set it in the .env file.")

# Create OpenAI client
openai_client = OpenAI()

# Entity extraction prompt template
ENTITY_EXTRACTION_TEMPLATE = """
You are a legal document analyzer specializing in entity extraction. 
Extract all entities from the following legal text and categorize them.

Text:
{text}

Extract the following entity types:
- Name: Person names (e.g., John Smith, Jane Doe)
- Date: Any dates (e.g., January 1, 2024, 01/01/2024)
- MonetaryAmount: Any monetary values (e.g., $500,000, 1.5 million dollars)
- LegalTerm: Legal terminology (e.g., trustee, beneficiary, executor)
- LegalClause: References to specific clauses or sections (e.g., Article 7, Section 3.2)

Format each entity as a JSON object with "entity" and "type" fields, and return a list of these objects.
For example: [{"entity": "John Smith", "type": "Name"}, {"entity": "January 1, 2024", "type": "Date"}]

IMPORTANT: Only use simple JSON format with keys "entity" and "type" without quotes in the keys.
"""

# Document generation prompt template
DOCUMENT_GENERATION_TEMPLATE = """
You are a legal document generator specializing in creating professional legal documents.
Create a {document_type} based on the following entities:

{entities_json}

Create a formal, well-structured legal document that:
1. Incorporates all the provided entities in appropriate context
2. Follows standard legal document formatting
3. Uses proper legal terminology
4. Includes appropriate headers, sections, and clauses
5. Maintains a professional tone throughout

The document should be complete and ready for review by legal professionals.
"""

class LLMProcessor:
    """
    LLM-based processor for enhanced entity extraction from legal documents.
    """
    
    def __init__(self, model_name="gpt-3.5-turbo"):
        """
        Initialize the LLM processor with the specified model.
        
        Args:
            model_name: The OpenAI model to use. Default is "gpt-3.5-turbo"
        """
        self.llm = ChatOpenAI(model_name=model_name, temperature=0)
        
        # Entity extraction chain - using text as the input variable
        self.entity_prompt = PromptTemplate(
            input_variables=["text"],
            template=ENTITY_EXTRACTION_TEMPLATE
        )
        self.entity_chain = LLMChain(
            llm=self.llm,
            prompt=self.entity_prompt
        )
        
        # Document generation chain
        self.document_llm = ChatOpenAI(model_name=model_name, temperature=0.2)
        self.document_prompt = PromptTemplate(
            input_variables=["document_type", "entities_json"],
            template=DOCUMENT_GENERATION_TEMPLATE
        )
        self.document_chain = LLMChain(
            llm=self.document_llm,
            prompt=self.document_prompt
        )
    
    def extract_entities(self, text: str) -> List[Dict[str, str]]:
        """
        Extract entities from the given text using the LLM.
        
        Args:
            text: The text to extract entities from
            
        Returns:
            A list of dictionaries, each containing an entity and its type
        """
        try:
            # Direct approach using the OpenAI client directly
            completion = openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a legal document analyzer specializing in entity extraction."},
                    {"role": "user", "content": f"Extract all entities from this legal text and categorize them:\n\n{text}\n\nExtract the following entity types:\n- Name: Person names\n- Date: Any dates\n- MonetaryAmount: Any monetary values\n- LegalTerm: Legal terminology\n- LegalClause: References to specific clauses\n\nFormat the response as a simple JSON array with each entity having 'entity' and 'type' fields. Example: [{{'entity': 'John Smith', 'type': 'Name'}}, {{'entity': 'January 1, 2024', 'type': 'Date'}}]"}
                ],
                temperature=0
            )
            
            # Extract the content from the response
            result = completion.choices[0].message.content
            
            # Find and extract the JSON list from the response
            start_idx = result.find('[')
            end_idx = result.rfind(']') + 1
            
            if start_idx >= 0 and end_idx > start_idx:
                json_str = result[start_idx:end_idx]
                try:
                    # Parse the JSON directly
                    entities = json.loads(json_str)
                    return entities
                except json.JSONDecodeError as e:
                    print(f"Error parsing JSON: {e}")
                    return []
            else:
                print("Could not find JSON output in LLM response")
                return []
        except Exception as e:
            print(f"Error extracting entities: {e}")
            return []
    
    def analyze_document(self, text: str) -> Dict:
        """
        Perform a comprehensive analysis of a legal document.
        
        Args:
            text: The document text to analyze
            
        Returns:
            A dictionary containing entities and other analysis results
        """
        entities = self.extract_entities(text)
        
        # Group entities by type
        entities_by_type = {}
        for entity in entities:
            entity_type = entity.get('type')
            if entity_type:
                if entity_type not in entities_by_type:
                    entities_by_type[entity_type] = []
                entities_by_type[entity_type].append(entity.get('entity'))
        
        return {
            "entities": entities,
            "entities_by_type": entities_by_type,
            "total_entities": len(entities)
        }
    
    def generate_document(self, entities: List[Dict[str, str]], document_type: str = "Trust Agreement") -> str:
        """
        Generate a legal document based on extracted entities.
        
        Args:
            entities: List of entity dictionaries, each with 'entity' and 'type' keys
            document_type: Type of document to generate (e.g., "Trust Agreement", "Will")
            
        Returns:
            Generated document text
        """
        try:
            # Convert entities to JSON string for the prompt
            entities_json = json.dumps(entities, indent=2)
            
            # Generate document
            response = self.document_chain.invoke({
                "document_type": document_type,
                "entities_json": entities_json
            })
            
            # Extract document text from response
            document_text = response.get('text', '')
            return document_text
        except Exception as e:
            print(f"Error generating document: {e}")
            return f"Error generating document: {str(e)}"
    
    def save_txt_document(self, document_text: str, filename: str = "generated_document.txt") -> str:
        """
        Save the generated document to a text file.
        
        Args:
            document_text: The document text to save
            filename: Name of the file to save the document to
            
        Returns:
            Path to the saved document
        """
        try:
            with open(filename, 'w') as f:
                f.write(document_text)
            return os.path.abspath(filename)
        except Exception as e:
            print(f"Error saving document: {e}")
            return ""
    
    def save_docx_document(self, document_text: str, filename: str = "generated_document.docx") -> str:
        """
        Save the generated document to a DOCX file.
        
        Args:
            document_text: The document text to save
            filename: Name of the file to save the document to
            
        Returns:
            Path to the saved document
        """
        try:
            # Create document
            doc = docx.Document()
            
            # Set basic styles
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # Split text into lines and process
            lines = document_text.strip().split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                
                # Simple formatting based on line characteristics
                if line.isupper():
                    # Likely a header
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.style = 'Heading 1'
                elif line.startswith('ARTICLE') or line.startswith('SECTION'):
                    p = doc.add_paragraph(line)
                    p.style = 'Heading 2'
                else:
                    p = doc.add_paragraph(line)
            
            # Save the document
            doc.save(filename)
            return os.path.abspath(filename)
        except Exception as e:
            print(f"Error creating DOCX: {e}")
            return ""


# Example usage
if __name__ == "__main__":
    # Sample text
    sample_text = """
    John Smith appointed Jane Doe as the trustee of his estate on March 15, 2024.
    The trust agreement allocates $500,000 to the beneficiary Sarah Williams.
    Upon the death of Thomas Anderson, the successor trustee Elizabeth Wilson 
    shall distribute $2.5 million according to Article 7.
    """
    
    processor = LLMProcessor()
    
    # Extract entities
    entities = processor.extract_entities(sample_text)
    print("\nExtracted Entities:")
    print("=" * 50)
    for entity in entities:
        print(f"Entity: {entity.get('entity')}, Type: {entity.get('type')}")
    
    # Generate document from entities
    document = processor.generate_document(entities)
    print("\nGenerated Document:")
    print("=" * 50)
    print(document)
    
    # Save document in both formats
    txt_path = processor.save_txt_document(document, "generated_document.txt")
    if txt_path:
        print(f"\nText document saved to: {txt_path}")
    
    docx_path = processor.save_docx_document(document, "generated_document.docx")
    if docx_path:
        print(f"\nDOCX document saved to: {docx_path}") 